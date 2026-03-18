#!/usr/bin/env python3
"""Generate Word document with full PPT content"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('皇庭国际中心数字化营销服务建议书', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('PPT制作内容大纲（共30页）')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# Page 1
doc.add_heading('第1页：封面', level=1)
doc.add_paragraph('主标题：皇庭国际中心数字化营销服务建议书')
doc.add_paragraph('副标题：打造CBD商务新地标 · 数字化营销全案解决方案')
doc.add_paragraph('提案方：XX营销服务机构')
doc.add_paragraph('日期：2024年X月')

doc.add_page_break()

# Page 2
doc.add_heading('第2页：目录', level=1)
doc.add_paragraph('1. 项目概况与定位')
doc.add_paragraph('2. 市场环境分析')
doc.add_paragraph('3. 数字化营销策略')
doc.add_paragraph('4. 第一太平戴维斯渠道策略')
doc.add_paragraph('5. 推广执行计划')
doc.add_paragraph('6. 预期效果与目标')
doc.add_paragraph('7. 合作方案')

doc.add_page_break()

# Page 3
doc.add_heading('第3页：项目概况——皇庭国际中心项目介绍', level=1)
doc.add_heading('【项目基本信息】', level=2)
doc.add_paragraph('皇庭国际中心位于城市CBD核心区，是皇庭集团倾力打造的旗舰商务综合体项目。项目总建筑面积12万平方米，其中地上38层，地下3层，业态包括甲级写字楼8万平方米、商业配套2万平方米、停车场2万平方米。项目预计2024年第三季度正式交付使用，定位为5A甲级智能商务办公空间。')

doc.add_heading('【核心数据】', level=2)
table = doc.add_table(rows=8, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = '指标'
table.rows[0].cells[1].text = '数据'
table.rows[1].cells[0].text = '总建筑面积'
table.rows[1].cells[1].text = '12万㎡'
table.rows[2].cells[0].text = '写字楼面积'
table.rows[2].cells[1].text = '8万㎡'
table.rows[3].cells[0].text = '商业配套'
table.rows[3].cells[1].text = '2万㎡'
table.rows[4].cells[0].text = '楼层'
table.rows[4].cells[1].text = '地上38层，地下3层'
table.rows[5].cells[0].text = '单层面积'
table.rows[5].cells[1].text = '1500-2000㎡'
table.rows[6].cells[0].text = '层高'
table.rows[6].cells[1].text = '4.2米（市场平均3.8米）'
table.rows[7].cells[0].text = '认证'
table.rows[7].cells[1].text = 'LEED金级预认证'

doc.add_heading('【建筑品质优势】', level=2)
doc.add_paragraph('项目由国际知名建筑设计事务所操刀设计，采用全玻璃幕墙配合LOW-E节能设计。配备智能化楼宇管理系统：人脸识别门禁系统、智能停车引导系统、5G网络全覆盖、楼宇自控系统。')

doc.add_heading('【区位交通优势】', level=2)
doc.add_paragraph('双地铁交汇上盖：距离地铁1号线、3号线交汇站仅200米，步行3分钟。30分钟可达国际机场，15分钟可达高铁站。周边配套：五星级酒店8家、大型购物中心6个、米其林星级餐厅12家。')

doc.add_heading('【物业服务】', level=2)
doc.add_paragraph('第一太平戴维斯提供24小时管家式国际化物业服务。')

doc.add_page_break()

# Page 4
doc.add_heading('第4页：目标客群定位——精准锁定高价值客户', level=1)
doc.add_heading('【三大目标客群】', level=2)

doc.add_heading('一、金融机构（占比35%）', level=3)
doc.add_paragraph('包括：银行分行、证券公司营业部、保险公司区域中心、基金公司')
doc.add_paragraph('特点：对办公环境形象要求高，注重楼宇品质，租金支付能力强')

doc.add_heading('二、科技企业（占比25%）', level=3)
doc.add_paragraph('包括：互联网公司、软件/AI企业、大数据企业')
doc.add_paragraph('特点：员工年轻化，注重智能化，追求办公舒适度')

doc.add_heading('三、专业服务机构（占比20%）', level=3)
doc.add_paragraph('包括：律师事务所、会计师事务所、咨询公司、广告公司')
doc.add_paragraph('特点：商务接待需求高，重视周边配套')

doc.add_heading('【客户典型特征】', level=2)
doc.add_paragraph('年营业收入5000万元以上，员工规模100-500人，日租金预算8-12元/㎡，选址决策周期3-6个月，决策人：CEO、COO或行政总监。')

doc.add_page_break()

# Page 5
doc.add_heading('第5页：市场环境分析——CBD商务区深度洞察', level=1)
doc.add_heading('【区域商务环境】', level=2)
doc.add_paragraph('城市第一CBD核心区，占地8平方公里。经济贡献：全市25%的GDP，30%的税收收入。产业聚集：60%金融机构总部 + 40%科技公司总部 + 50%专业服务机构。')

doc.add_heading('【市场需求分析】', level=2)
doc.add_paragraph('现有甲级写字楼200万㎡，入驻企业5000+家，就业人口30万。约30%现有企业有升级办公环境需求，预计150家企业将搬迁/扩租。每年新增注册企业2000家，约10%符合甲级写字楼标准。')

doc.add_heading('【竞争格局与机会】', level=2)
doc.add_paragraph('区域内甲级写字楼项目15个，总体空置率8%（供不应求）。未来12个月内无新增供应，皇庭国际中心拥有宝贵的时间窗口。')

doc.add_page_break()

# Page 6
doc.add_heading('第6页：竞争项目对比分析', level=1)
doc.add_heading('【竞争项目A】位于项目东侧800米，2019年入市', level=2)
doc.add_paragraph('基本情况：10万㎡，出租率85%，日租金9-11元/㎡')
doc.add_paragraph('优势：品牌知名度高，已入驻多家知名企业')
doc.add_paragraph('劣势：楼龄5年，设施老化，智能化程度低')

doc.add_heading('【竞争项目B】位于项目西侧1.2公里，2021年入市', level=2)
doc.add_paragraph('基本情况：8万㎡，出租率78%，日租金8-10元/㎡')
doc.add_paragraph('优势：租金定价较低')
doc.add_paragraph('劣势：距地铁站远（步行15分钟），周边配套不完善')

doc.add_heading('【皇庭国际中心核心优势】', level=2)
doc.add_paragraph('交通便利性最佳：双地铁交汇上盖，距地铁站仅200米')
doc.add_paragraph('建筑品质最新：2024年全新交付，智能化程度最高')
doc.add_paragraph('物业服务最优：第一太平戴维斯国际化标准服务')
doc.add_paragraph('目标租金10-12元/㎡，综合性价比高')

doc.add_page_break()

# Page 7
doc.add_heading('第7页：数字化营销策略（章节页）', level=1)
doc.add_paragraph('章节标题：数字化营销策略')
doc.add_paragraph('副标题：全渠道数字化营销体系，精准触达目标客户')

doc.add_page_break()

# Page 8
doc.add_heading('第8页：官方网站与小程序生态建设', level=1)
doc.add_heading('【网站建设目标】', level=2)
doc.add_paragraph('打造高品质响应式官网，适配PC/平板/手机多端。采用最新前端技术，页面加载速度控制在3秒以内。')

doc.add_heading('【六大核心功能模块】', level=2)
doc.add_paragraph('1. VR全景看房系统——360度全景拍摄，身临其境感受项目品质')
doc.add_paragraph('2. 楼层平面展示功能——查看各楼层详细平面图，了解户型分布')
doc.add_paragraph('3. 实时房源查询系统——实时显示可租面积、租金价格、朝向景观')
doc.add_paragraph('4. 在线预约看房功能——填写信息、选择看房时间，自动发送确认')
doc.add_paragraph('5. 智能客服机器人（ChatGPT技术）——7×24小时在线解答，自动转接人工')
doc.add_paragraph('6. 数据分析系统——埋点追踪用户行为，优化营销策略')

doc.add_page_break()

# Page 9
doc.add_heading('第9页：短视频与直播营销矩阵', level=1)
doc.add_heading('【短视频平台布局】', level=2)
doc.add_paragraph('抖音企业号 + 视频号 + 小红书企业号，三大平台同步运营，形成矩阵化营销布局。')

doc.add_heading('【内容规划】', level=2)
doc.add_paragraph('核心主题："高端商务办公生活方式"')
doc.add_paragraph('内容类型：项目宣传片、入驻企业访谈、办公环境展示、周边配套探店')

doc.add_heading('【发布运营策略】', level=2)
doc.add_paragraph('发布频率：每周3-5条')
doc.add_paragraph('付费投放：抖音DOU+精准定向25-45岁企业管理人员')

doc.add_heading('【直播看房】', level=2)
doc.add_paragraph('每周1-2场直播，专业置业顾问实地带领参观，设置直播专属优惠促进转化。')

doc.add_page_break()

# Page 10
doc.add_heading('第10页：精准数字广告投放策略', level=1)
doc.add_heading('【微信朋友圈广告】', level=2)
doc.add_paragraph('月均投放预算：15-20万元')
doc.add_paragraph('定向策略：年龄28-50岁，职位总监级以上，行业金融/科技/专业服务，地域项目周边10公里')

doc.add_heading('【百度系广告投放】', level=2)
doc.add_paragraph('搜索广告：购买"写字楼出租"、"CBD写字楼"等核心关键词')
doc.add_paragraph('信息流广告：百度APP、百家号、好看视频精准推送')

doc.add_heading('【LinkedIn广告投放】', level=2)
doc.add_paragraph('针对外企和跨国公司，制作英文版项目介绍，定向投放给跨国企业高管。')

doc.add_page_break()

# Page 11
doc.add_heading('第11页：内容营销与私域流量运营', level=1)
doc.add_heading('【内容营销】', level=2)
doc.add_paragraph('行业报告：《2024城市写字楼市场白皮书》、《CBD商务区企业选址趋势报告》、《未来办公空间设计趋势研究》')
doc.add_paragraph('专家IP打造：首席运营官个人IP，定期输出专业观点')
doc.add_paragraph('客户案例包装：制作入驻企业成功案例视频')

doc.add_heading('【私域流量运营】', level=2)
doc.add_paragraph('企业微信客户池：标签化管理，按行业/面积/预算分层运营')
doc.add_paragraph('皇庭企业家俱乐部：定期举办线上分享会、政策解读沙龙')

doc.add_page_break()

# Page 12
doc.add_heading('第12页：第一太平戴维斯渠道策略（章节页）', level=1)
doc.add_paragraph('章节标题：第一太平戴维斯渠道策略')
doc.add_paragraph('副标题：依托全球网络资源，打造高效招商渠道体系')

doc.add_page_break()

# Page 13
doc.add_heading('第13页：第一太平戴维斯资源价值分析', level=1)
doc.add_heading('【全球领先的房地产服务商】', level=2)
doc.add_paragraph('成立时间：1855年（160年历史）')
doc.add_paragraph('全球办事处：600多个，覆盖70多个国家')
doc.add_paragraph('中国分公司：15个主要城市，3000+员工')
doc.add_paragraph('服务客户：苹果、谷歌、微软、摩根大通等500+世界500强')

doc.add_heading('【核心价值】', level=2)
doc.add_paragraph('客户数据库：超过500家跨国企业中国区总部或区域中心')
doc.add_paragraph('全球网络：帮助接触尚未进入中国市场的国际企业')
doc.add_paragraph('本地团队：每个城市50-80名专业商业地产顾问')

doc.add_page_break()

# Page 14
doc.add_heading('第14页：渠道合作体系构建', level=1)
doc.add_heading('【独家代理权】', level=2)
doc.add_paragraph('授予第一太平戴维斯项目周边3公里范围内的独家代理权')

doc.add_heading('【其他渠道布局】', level=2)
doc.add_paragraph('国际五大行：仲量联行、世邦魏理仕、戴德梁行、高力国际')
doc.add_paragraph('本地中介：筛选20-30家优质中介公司')

doc.add_heading('【差异化佣金政策】', level=2)
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = '渠道类型'
table.rows[0].cells[1].text = '佣金标准'
table.rows[0].cells[2].text = '额外奖励'
table.rows[1].cells[0].text = '第一太平戴维斯'
table.rows[1].cells[1].text = '首年租金2个月'
table.rows[1].cells[2].text = '业绩达标额外0.5个月'
table.rows[2].cells[0].text = '其他五大行'
table.rows[2].cells[1].text = '首年租金1.5个月'
table.rows[2].cells[2].text = '无'
table.rows[3].cells[0].text = '本地中介'
table.rows[3].cells[1].text = '首年租金1个月'
table.rows[3].cells[2].text = '无'

doc.add_page_break()

# Page 15
doc.add_heading('第15页：客户来源渠道占比规划', level=1)
doc.add_heading('【渠道贡献目标预测】', level=2)
table = doc.add_table(rows=5, cols=4)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = '渠道'
table.rows[0].cells[1].text = '占比'
table.rows[0].cells[2].text = '客户数'
table.rows[0].cells[3].text = '客户特征'
table.rows[1].cells[0].text = '第一太平戴维斯'
table.rows[1].cells[1].text = '40%'
table.rows[1].cells[2].text = '约16家'
table.rows[1].cells[3].text = '跨国企业和国内大型企业，单客面积大'
table.rows[2].cells[0].text = '其他五大行'
table.rows[2].cells[1].text = '25%'
table.rows[2].cells[2].text = '约10家'
table.rows[2].cells[3].text = '国内大型企业客户'
table.rows[3].cells[0].text = '本地中介'
table.rows[3].cells[1].text = '20%'
table.rows[3].cells[2].text = '约8家'
table.rows[3].cells[3].text = '中小企业客户'
table.rows[4].cells[0].text = '自主获客'
table.rows[4].cells[1].text = '15%'
table.rows[4].cells[2].text = '约6家'
table.rows[4].cells[3].text = '数字化渠道直接获取'

doc.add_page_break()

# Page 16
doc.add_heading('第16页：推广执行计划（章节页）', level=1)
doc.add_paragraph('章节标题：推广执行计划')
doc.add_paragraph('副标题：分阶段推进，确保营销目标达成')

doc.add_page_break()

# Page 17
doc.add_heading('第17页：第一阶段——数字化基建与渠道搭建', level=1)
doc.add_heading('【时间】第1-2个月', level=2)
doc.add_heading('【核心任务】', level=2)
doc.add_paragraph('完成所有数字化营销基础设施的搭建和渠道合作体系的建立。')
doc.add_heading('【具体工作】', level=2)
doc.add_paragraph('1. 官网和小程序开发——响应式官网，包含VR看房、在线预约、智能客服')
doc.add_paragraph('2. 短视频账号搭建——注册认证抖音、视频号、小红书，制作首批20条短视频')
doc.add_paragraph('3. 广告投放准备——完成微信、百度、LinkedIn广告开户和资质审核')
doc.add_paragraph('4. 渠道合作签约——与第一太平戴维斯完成独家代理协议签署')
doc.add_heading('【阶段目标】', level=2)
doc.add_paragraph('官网上线日均500访问量，短视频粉丝5000+，渠道体系覆盖50+名专业顾问')

doc.add_page_break()

# Page 18
doc.add_heading('第18页：第二阶段——全面推广与客户积累', level=1)
doc.add_heading('【时间】第3-8个月', level=2)
doc.add_heading('【核心任务】', level=2)
doc.add_paragraph('营销推广的核心期，全面启动各项推广活动，大量积累客户资源。')
doc.add_heading('【具体工作】', level=2)
doc.add_paragraph('1. 数字广告大规模投放——微信朋友圈月投15-20万，百度5-8万，LinkedIn 3-5万，DOU+ 5万')
doc.add_paragraph('2. 短视频持续输出——每周3-5条，预计6个月粉丝达5万+')
doc.add_paragraph('3. 直播看房常态化——每周1-2场直播，收集高意向客户')
doc.add_paragraph('4. 渠道带客全面铺开——每周定向拜访目标企业，每月1-2场渠道推介会')
doc.add_heading('【阶段目标】', level=2)
doc.add_paragraph('累计客户咨询3000+组，实地看房500+组，签约客户20+家，出租率达到60%')

doc.add_page_break()

# Page 19
doc.add_heading('第19页：第三阶段——精准转化与客户维护', level=1)
doc.add_heading('【时间】第9-12个月', level=2)
doc.add_heading('【核心任务】', level=2)
doc.add_paragraph('转化前期积累的意向客户，同时做好已签约客户的服务和口碑建设。')
doc.add_heading('【具体工作】', level=2)
doc.add_paragraph('1. 意向客户攻坚——A类客户项目总监亲自跟进，每周至少一次沟通')
doc.add_paragraph('2. 老客户转介绍——建立"老带新"奖励机制，成功推荐奖励1个月租金减免')
doc.add_paragraph('3. 私域流量运营——企业微信客户池标签化管理，皇庭企业家俱乐部定期活动')
doc.add_paragraph('4. 服务品质提升——满意度调研，快速响应机制48小时内解决问题')
doc.add_heading('【阶段目标】', level=2)
doc.add_paragraph('新增签约客户20+家，累计出租率达到80%，客户满意度90%以上，老客户转介绍率30%')

doc.add_page_break()

# Page 20
doc.add_heading('第20页：12个月出租率增长曲线', level=1)
doc.add_heading('【出租率增长路径】', level=2)
table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = '阶段'
table.rows[0].cells[1].text = '时间'
table.rows[0].cells[2].text = '出租率'
table.rows[1].cells[0].text = '基建期'
table.rows[1].cells[1].text = '第1-2个月'
table.rows[1].cells[2].text = '10%'
table.rows[2].cells[0].text = '快速推广期'
table.rows[2].cells[1].text = '第3-6个月'
table.rows[2].cells[2].text = '35%→60%'
table.rows[3].cells[0].text = '持续推广期'
table.rows[3].cells[1].text = '第7-9个月'
table.rows[3].cells[2].text = '75%'
table.rows[4].cells[0].text = '收尾期'
table.rows[4].cells[1].text = '第10-12个月'
table.rows[4].cells[2].text = '80%'

doc.add_paragraph()
doc.add_paragraph('增长曲线数据点：Month 1: 10% → Month 3: 35% → Month 6: 60% → Month 9: 75% → Month 12: 80%')

doc.add_page_break()

# Page 21
doc.add_heading('第21页：预期效果与目标（章节页）', level=1)
doc.add_paragraph('章节标题：预期效果与目标')
doc.add_paragraph('副标题：量化目标，确保可衡量、可追踪')

doc.add_page_break()

# Page 22
doc.add_heading('第22页：数字化营销效果目标', level=1)
doc.add_heading('【12个月营销目标】', level=2)
doc.add_paragraph('官网流量：累计访问量10万+，月均8000+（搜索30% + 广告50% + 社交20%）')
doc.add_paragraph('小程序用户：累计5000+，活跃用户40%，预约看房800+组')
doc.add_paragraph('短视频营销：粉丝8万+，播放量500万+，引流3000+人次')
doc.add_paragraph('直播看房：40+场，观看10万+人次，收集高意向客户500+组')
doc.add_paragraph('私域流量：企业微信2000+人，VIP社群300+人，转化率15%')

doc.add_page_break()

# Page 23
doc.add_heading('第23页：招商效果与投资回报目标', level=1)
doc.add_heading('【总体招商目标】', level=2)
doc.add_paragraph('首年累计签约客户40家以上，累计出租面积6.4万㎡，出租率达到80%')
doc.add_paragraph('客户结构：龙头企业3-5家（1500㎡+），中型客户15-20家（500-1500㎡），小型客户15-20家（300-500㎡）')
doc.add_heading('【经济效益】', level=2)
doc.add_paragraph('首年租金收入：约2亿元（按日租金10元/㎡、出租率70%计算）')
doc.add_paragraph('投资回报周期：预计3-4年，第5年起年租金收入可达2.5亿元')
doc.add_heading('【品牌价值】', level=2)
doc.add_paragraph('打造区域标杆性甲级写字楼项目，获得"年度最佳商务楼宇"等行业奖项')
doc.add_heading('【客户口碑】', level=2)
doc.add_paragraph('客户满意度90%以上，续约率85%以上，老客户每年转介绍10-15组')

doc.add_page_break()

# Page 24
doc.add_heading('第24页：合作方案（章节页）', level=1)
doc.add_paragraph('章节标题：合作方案')
doc.add_paragraph('副标题：灵活的合作模式，共赢的伙伴关系')

doc.add_page_break()

# Page 25
doc.add_heading('第25页：服务内容与费用结构', level=1)
doc.add_heading('【全案营销代理服务】', level=2)
doc.add_paragraph('涵盖数字化营销、渠道管理、客户对接、签约转化等全流程服务。')
doc.add_heading('【费用结构】', level=2)
table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = '费用类型'
table.rows[0].cells[1].text = '金额'
table.rows[0].cells[2].text = '说明'
table.rows[1].cells[0].text = '基础服务费'
table.rows[1].cells[1].text = '80万元/年'
table.rows[1].cells[2].text = '团队人力+内容制作+活动执行'
table.rows[2].cells[0].text = '数字广告预算'
table.rows[2].cells[1].text = '150-200万元/年'
table.rows[2].cells[2].text = '按实际投放结算'
table.rows[3].cells[0].text = '第一太平戴维斯佣金'
table.rows[3].cells[1].text = '首年租金2个月'
table.rows[3].cells[2].text = '独家代理'
table.rows[4].cells[0].text = '其他渠道佣金'
table.rows[4].cells[1].text = '首年租金1.5个月'
table.rows[4].cells[2].text = '按成交结算'

doc.add_heading('【项目团队配置】', level=2)
doc.add_paragraph('项目总监1名 + 数字营销经理1名 + 渠道经理1名 + 内容策划1名 + 客户顾问2名')

doc.add_page_break()

# Page 26
doc.add_heading('第26页：第一太平戴维斯专属合作条款', level=1)
doc.add_heading('【独家代理权】', level=2)
doc.add_paragraph('授予项目周边3公里范围内独家代理权，有效期12个月，业绩达标可优先续约。')
doc.add_heading('【全球客户资源优先推荐】', level=2)
doc.add_paragraph('承诺将全球客户数据库中计划在中国设立办公室的企业客户优先推荐，享有全球客户推介会优先参与权。')
doc.add_heading('【专属服务团队】', level=2)
doc.add_paragraph('组建不少于5人的专门服务团队，包括1名总监级负责人和4名资深顾问，每周至少驻场2天。')
doc.add_heading('【佣金保障与额外奖励】', level=2)
doc.add_paragraph('享受最高佣金标准（首年租金2个月），业绩达标额外0.5个月奖励，世界500强企业每签约1家额外奖励5万元。')

doc.add_page_break()

# Page 27
doc.add_heading('第27页：双方职责分工', level=1)
doc.add_heading('【我方（营销服务商）职责】', level=2)
doc.add_paragraph('1. 制定整体营销策略和执行方案，每月提交执行计划和上月总结报告')
doc.add_paragraph('2. 负责数字化营销的全面执行，包括官网运营、广告投放、短视频制作、直播活动')
doc.add_paragraph('3. 负责渠道管理和维护，与第一太平戴维斯等渠道保持密切沟通，组织渠道培训')
doc.add_paragraph('4. 提供客户对接和商务谈判支持，协助客户看房、方案定制、合同谈判')
doc.add_paragraph('5. 定期组织客户活动，如项目品鉴会、行业论坛、客户答谢宴')
doc.add_paragraph('6. 提供全程数据分析和效果追踪，每周提交数据分析报告')
doc.add_heading('【贵方（项目方）职责】', level=2)
doc.add_paragraph('1. 提供完整的项目资料，包括建筑图纸、技术参数、交付标准、政策文件')
doc.add_paragraph('2. 配合客户带看和商务谈判，安排项目负责人解答客户技术问题')
doc.add_paragraph('3. 审批营销方案、广告创意、活动方案，确保符合公司品牌要求')
doc.add_paragraph('4. 按时支付服务费用、广告费用和渠道佣金')
doc.add_paragraph('5. 提供必要的场地支持，如举办活动所需的会议室、展示区')
doc.add_paragraph('6. 指定专人负责日常对接，及时响应我们的需求和建议')

doc.add_page_break()

# Page 28
doc.add_heading('第28页：服务团队介绍', level=1)
doc.add_heading('【项目总监】', level=2)
doc.add_paragraph('张经理，15年商业地产从业经验，曾成功操盘3个甲级写字楼项目，累计招商面积超过50万平方米。')
doc.add_heading('【数字营销经理】', level=2)
doc.add_paragraph('李经理，8年互联网营销经验，曾任职知名互联网公司，擅长短视频运营、信息流广告投放、私域流量运营。')
doc.add_heading('【渠道经理】', level=2)
doc.add_paragraph('王经理，10年商业地产渠道管理经验，与五大行及本地中介建立广泛联系网络，擅长渠道激励和关系维护。')
doc.add_heading('【内容策划】', level=2)
doc.add_paragraph('陈经理，6年商业地产内容创作经验，擅长文案撰写、视频策划、品牌包装，多条短视频播放量超过百万。')

doc.add_page_break()

# Page 29
doc.add_heading('第29页：下一步行动计划', level=1)
doc.add_heading('【时间轴】', level=2)
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = '时间'
table.rows[0].cells[1].text = '行动'
table.rows[1].cells[0].text = '本周内'
table.rows[1].cells[1].text = '安排项目实地考察，深入了解项目情况'
table.rows[2].cells[0].text = '下周内'
table.rows[2].cells[1].text = '完善营销方案细节，提交执行计划和预算表'
table.rows[3].cells[0].text = '两周内'
table.rows[3].cells[1].text = '签署正式合作协议，组建服务团队'
table.rows[4].cells[0].text = '一个月内'
table.rows[4].cells[1].text = '官网小程序上线，举办首场项目品鉴会'

doc.add_page_break()

# Page 30
doc.add_heading('第30页：感谢聆听', level=1)
doc.add_paragraph('期待与您携手，共创皇庭国际中心辉煌！')
doc.add_paragraph()
doc.add_heading('【联系我们】', level=2)
doc.add_paragraph('项目总监：张经理')
doc.add_paragraph('联系电话：138-XXXX-XXXX')
doc.add_paragraph('电子邮箱：zhang@company.com')
doc.add_paragraph('公司地址：XX市XX区XX路XX号')

# Save
doc.save('/Users/bob/.openclaw/workspace/皇庭国际中心PPT内容大纲.docx')
print("Word document generated: 皇庭国际中心PPT内容大纲.docx")
